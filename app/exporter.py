import json
import csv
import logging
from typing import List, Dict, Any
from app.indexer.sqlite_store import SQLiteStore

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class Exporter:
    """Export summaries and results with decryption."""

    def __init__(self, sqlite_store: SQLiteStore):
        self.sqlite_store = sqlite_store

    def export_summary(self, summaries: Dict[str, str], file_path: str, format: str = 'json'):
        """Export summaries to file."""
        try:
            # Decrypt texts if needed
            decrypted = {k: self.sqlite_store._decrypt_text(v) if isinstance(v, bytes) else v for k, v in summaries.items()}
            if format == 'json':
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(decrypted, f, ensure_ascii=False, indent=4)
            elif format == 'csv':
                with open(file_path, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.DictWriter(f, fieldnames=['chunk_id', 'summary'])
                    writer.writeheader()
                    for k, v in decrypted.items():
                        writer.writerow({'chunk_id': k, 'summary': v})
            logger.info(f"Exported to {file_path}")
        except Exception as e:
            logger.error(f"Export failed: {str(e)}")
            raise

    def load_exported(self, file_path: str, format: str = 'json') -> Dict[str, Any]:
        """Load exported file with error handling."""
        try:
            if format == 'json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            elif format == 'csv':
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.DictReader(f)
                    return {row['chunk_id']: row['summary'] for row in reader}
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error: {str(e)} - Check file encoding or corruption.")
            raise
        except Exception as e:
            logger.error(f"Load failed: {str(e)}")
            raise

if __name__ == "__main__":
    # Test code
    pass
import os
import docx
from fpdf import FPDF
from typing import List, Dict, Optional
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class Exporter:
    def __init__(self, output_dir: str):
        """Initialize exporter with output directory."""
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def export_to_txt(self, content: List[Dict[str, str]], output_filename: str) -> bool:
        """Export content to TXT file."""
        try:
            output_path = os.path.join(self.output_dir, output_filename)
            with open(output_path, 'w', encoding='utf-8') as f:
                for item in content:
                    f.write(f"{item.get('text', '')}\n")
                    if 'tags' in item:
                        f.write(f"Tags: {', '.join(item['tags'])}\n")
                    f.write("\n")
            logger.info(f"Exported to TXT: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to export to TXT: {e}")
            return False

    def export_to_docx(self, content: List[Dict[str, str]], output_filename: str) -> bool:
        """Export content to DOCX file."""
        try:
            doc = docx.Document()
            for item in content:
                doc.add_paragraph(item.get('text', ''))
                if 'tags' in item:
                    doc.add_paragraph(f"Tags: {', '.join(item['tags'])}")
                doc.add_paragraph("")
            output_path = os.path.join(self.output_dir, output_filename)
            doc.save(output_path)
            logger.info(f"Exported to DOCX: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to export to DOCX: {e}")
            return False

    def export_to_pdf(self, content: List[Dict[str, str]], output_filename: str) -> bool:
        """Export content to PDF file."""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.add_font('NotoSerifBengali', '', 'NotoSerifBengali-Regular.ttf', uni=True)
            pdf.set_font('NotoSerifBengali', '', 12)
            for item in content:
                pdf.multi_cell(0, 10, item.get('text', '').encode('utf-8').decode('utf-8'))
                if 'tags' in item:
                    pdf.multi_cell(0, 10, f"Tags: {', '.join(item['tags'])}".encode('utf-8').decode('utf-8'))
                pdf.ln(10)
            output_path = os.path.join(self.output_dir, output_filename)
            pdf.output(output_path)
            logger.info(f"Exported to PDF: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to export to PDF: {e}")
            return False

if __name__ == "__main__":
    # Test the exporter
    test_content = [
        {"text": "This is a summary.", "tags": ["summary", "test"]},
        {"text": "Another summary.", "tags": ["example", "doc"]}
    ]
    exporter = Exporter("output")
    exporter.export_to_txt(test_content, "test_summary.txt")
    exporter.export_to_docx(test_content, "test_summary.docx")
    exporter.export_to_pdf(test_content, "test_summary.pdf")