import docx
import os
from typing import List, Dict, Optional
from pathlib import Path
import sys
import logging

try:
    # Prefer to use the Normalizer class from app.normalizer
    from app.normalizer import Normalizer
except Exception:
    from normalizer import Normalizer

# Instantiate a module-level normalizer with defaults
normalizer = Normalizer()

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class DocxExtractor:
    """Simple DOCX extractor that returns paragraphs and table cells with basic metadata."""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = None
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        try:
            self.doc = docx.Document(file_path)
        except Exception as e:
            logger.error(f"Failed to open DOCX {file_path}: {e}")
            raise

    def extract_text(self) -> List[Dict[str, str]]:
        """Extract text segments from the DOCX file.

        Returns a list of dicts: {"text": str, "type": "paragraph"|"table", "style": str}
        """
        if not self.doc:
            raise ValueError("DOCX document not loaded")

        segments: List[Dict[str, str]] = []

        # Paragraphs
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            segments.append({
                "text": normalizer.normalize_text(text),
                "type": "paragraph",
                "style": getattr(para.style, "name", "")
            })

        # Tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if not text:
                        continue
                    segments.append({
                        "text": normalizer.normalize_text(text),
                        "type": "table",
                        "style": "table_cell"
                    })

        logger.info(f"Extracted {len(segments)} text segments from {self.file_path}")
        return segments

    def extract_metadata(self) -> Dict[str, str]:
        metadata: Dict[str, str] = {}
        try:
            props = self.doc.core_properties
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
            }
        except Exception:
            logger.debug("No core properties available or failed to read them")
        return metadata


def main():
    import argparse

    parser = argparse.ArgumentParser(description="DOCX extractor")
    parser.add_argument("docx_path", help="Path to .docx file")
    parser.add_argument("--output", "-o", help="Optional output text file")
    args = parser.parse_args()

    path = Path(args.docx_path)
    if not path.exists() or path.suffix.lower() != ".docx":
        logger.error("Please provide an existing .docx file")
        sys.exit(1)

    extractor = DocxExtractor(str(path))
    segments = extractor.extract_text()
    metadata = extractor.extract_metadata()

    print(f"Extracted {len(segments)} segments")
    print("Metadata:", metadata)

    if args.output:
        out = Path(args.output)
        with out.open("w", encoding="utf-8") as f:
            for s in segments:
                f.write(s["text"] + "\n\n")


if __name__ == "__main__":
    main()